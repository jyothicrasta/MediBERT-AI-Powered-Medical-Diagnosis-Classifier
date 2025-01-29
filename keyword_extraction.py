import os
import re
import pandas as pd
from docx import Document
from transformers import pipeline, BertTokenizer, BertForSequenceClassification
import streamlit as st
import torch
import joblib

# Load Hugging Face model for keyword extraction
keyphrase_extractor = pipeline(
    "ner",
    model="ml6team/keyphrase-extraction-distilbert-inspec",
    aggregation_strategy="simple"
)

# Functions for Keyword Extraction
def extract_keywords(text):
    try:
        keywords = keyphrase_extractor(text)
        tokens = [kw['word'] for kw in keywords]
        return merge_and_correct_keywords(tokens)
    except Exception as e:
        print(f"Keyword extraction failed: {e}")
        return []

def merge_and_correct_keywords(tokens):
    merged_keywords = []
    current_keyword = ""

    for token in tokens:
        if token.startswith("##"):
            current_keyword += token[2:]
        else:
            if current_keyword:
                merged_keywords.append(current_keyword)
            current_keyword = token
    if current_keyword:
        merged_keywords.append(current_keyword)

    return [dynamic_term_correction(kw.strip()) for kw in merged_keywords]

def dynamic_term_correction(term):
    corrections = {
        "ysplastic": "dysplastic",
        # Add other common corrections here
    }
    for error, correction in corrections.items():
        if error in term:
            return term.replace(error, correction)
    return term

# Functions for Text Extraction and Cleaning
def extract_text_from_docx(docx_file):
    try:
        doc = Document(docx_file)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)
    except Exception as e:
        print(f"Failed to extract text from {docx_file}: {e}")
        return ""

def remove_duplicates(text):
    lines = text.split('\n')
    unique_lines = list(dict.fromkeys([line.strip() for line in lines if line.strip()]))
    return "\n".join(unique_lines)

def clean_content(text):
    text = re.sub(r'…… ', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# Section Extraction
patterns = {
    'microscopic appearance': r'microscopic appearance\s*(.+?)\s+diagnosis',
    'gross description': r'gross description\s*(.+?)\s+microscopic appearance'
}

def extract_info(text, patterns):
    info = {}
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        info[key] = match.group(1).strip() if match else ""
    return info

# Categorization of Gross Description
def categorize_gross_description(text):
    import re

    categories = {
        "Number of pieces": [r"\b(one|two|three|four|five|\d+)\s+(soft tissue|hard tissue)s?\b"],
        "Size": [r"(\d+(\.\d+)?\s*x\s*\d+(\.\d+)?\s*x\s*\d+(\.\d+)?\s*cm)"],
        "Color": [r"\b(white|brown|grey|gray|red|blue|black|yellow)\b"],
        "Type of tissue": [
            r"\bsoft tissue\b", r"\bhard tissue\b", r"\bsoft tissues\b", r"\bhard tissues\b",
            r"\bsmall tissue\b", r"\blarge tissue\b"
        ]
    }

    results = {}
    for category, patterns in categories.items():
        matches = []
        for pattern in patterns:
            matches.extend(re.findall(pattern, text))
        # Flatten the matches and remove duplicates
        results[category] = list(set([match[0] if isinstance(match, tuple) else match for match in matches]))
     # Process "Type of tissue" category
    if "Type of tissue" in results:
        tissue_types = results["Type of tissue"]

        # Check if "soft tissue" or "hard tissue" is present
        if "soft tissue" in tissue_types or "hard tissue" in tissue_types:
            # If either or both are present, keep them as is and remove others
            results["Type of tissue"] = [
                item for item in tissue_types if item in ["soft tissue", "hard tissue"]
            ]
        else:
            # If neither "soft tissue" nor "hard tissue" is present,
            # replace "small tissue" and "large tissue" with "presence of tissue"
            results["Type of tissue"] = [
                "presence of tissue" if item in ["small tissue", "large tissue"] else item
                for item in tissue_types
            ]
            # Remove duplicates and keep only "presence of tissue" if applicable
            if "presence of tissue" in results["Type of tissue"]:
                results["Type of tissue"] = ["presence of tissue"]

    return results

# File Processing
def process_file(file_path):
    try:
        extracted_text = extract_text_from_docx(file_path)
        unique_text = remove_duplicates(extracted_text)
        cleaned_text = clean_content(unique_text)
        lower_text = cleaned_text.lower()
        sections = extract_info(lower_text, patterns)

        if sections.get("gross description"):
            sections["Gross Description Categories"] = categorize_gross_description(sections["gross description"])
        return sections
    except Exception as e:
        print(f"Failed to process file {file_path}: {e}")
        return None

# Load Model for Diagnosis Prediction
@st.cache_resource
def load_model_and_resources():
    model = BertForSequenceClassification.from_pretrained(r"D:\Health\TRAIN MODEL\Using Python\medical_diagnosis_bert_model")
    tokenizer = BertTokenizer.from_pretrained(r"D:\Health\TRAIN MODEL\Using Python\medical_diagnosis_bert_model")
    label_encoder = joblib.load(r"D:\Health\TRAIN MODEL\Using Python\label_encoder.pkl")
    return model, tokenizer, label_encoder

model, tokenizer, label_encoder = load_model_and_resources()

@st.cache_resource
def get_keywords():
    return joblib.load(r"D:\Health\TRAIN MODEL\Using Python\microscopic_keywords.pkl")

keywords_list = get_keywords()

def predict_diagnosis(keywords):
    input_text = " ".join(keywords)
    encoding = tokenizer(
        input_text,
        padding="max_length",
        truncation=True,
        max_length=128,
        return_tensors="pt"
    )
    with torch.no_grad():
        outputs = model(**{key: val.to(model.device) for key, val in encoding.items()})
    logits = outputs.logits
    predicted_class = torch.argmax(logits, dim=1).item()
    diagnosis = label_encoder.inverse_transform([predicted_class])[0]
    return diagnosis

# Streamlit App
def main():
    st.title("Medical Report Analysis Tool")

    page = st.sidebar.radio("Choose a task", ["Keyword Extraction", "Gross Description Categorization", "Diagnosis Prediction"])

    if page == "Keyword Extraction":
        st.subheader("Extract Keywords from Microscopic Appearance")
        uploaded_file = st.file_uploader("Upload a Word document", type="docx")

        if uploaded_file:
            with open("temp_upload.docx", "wb") as f:
                f.write(uploaded_file.getbuffer())
            file_result = process_file("temp_upload.docx")
            if file_result and "microscopic appearance" in file_result:
                keywords = extract_keywords(file_result["microscopic appearance"])
                st.write(pd.DataFrame({"Keywords": keywords}))
            else:
                st.error("No 'microscopic appearance' section found.")

    elif page == "Gross Description Categorization":
        st.subheader("Categorize Gross Description")
        uploaded_file = st.file_uploader("Upload a Word document", type="docx")

        if uploaded_file:
            with open("temp_upload.docx", "wb") as f:
                f.write(uploaded_file.getbuffer())
            file_result = process_file("temp_upload.docx")
            if file_result and "gross description" in file_result:
                categories = file_result.get("Gross Description Categories", {})
                st.write(pd.DataFrame(categories.items(), columns=["Feature", "Category"]))
            else:
                st.error("No 'gross description' section found.")

    elif page == "Diagnosis Prediction":
        st.subheader("Predict Diagnosis from Keywords")
        selected_keywords = st.multiselect("Select keywords:", keywords_list)

        if st.button("Predict Diagnosis"):
            if selected_keywords:
                diagnosis = predict_diagnosis(selected_keywords)
                st.success(f"Predicted Diagnosis: {diagnosis}")
            else:
                st.warning("Please select at least one keyword.")

if __name__ == "__main__":
    main()
